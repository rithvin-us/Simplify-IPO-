"""POST /export — render approved sections to a DRHP document.

Always returns markdown + printable HTML (browser -> PDF). Also returns a
base64 .docx when python-docx is installed. Includes a compliance summary.
"""
from __future__ import annotations

import base64
import io
from datetime import date

from fastapi import APIRouter
from pydantic import BaseModel

from ..schema import SECTION_KEYS

router = APIRouter(tags=["export"])


class ExportSection(BaseModel):
    section_key: str
    title: str
    content: str
    status: str = ""


class ExportRequest(BaseModel):
    company_name: str = "The Company"
    sections: list[ExportSection] = []


def _compliance_summary(req: ExportRequest) -> dict:
    present = {s.section_key for s in req.sections if s.content.strip()}
    missing_sections = [k for k in SECTION_KEYS if k not in present]
    info_required = sum(s.content.count("[INFORMATION REQUIRED") for s in req.sections)
    finalized = sum(1 for s in req.sections if s.status in ("final", "legal_reviewed"))
    return {
        "sections_total": len(SECTION_KEYS),
        "sections_present": len(present),
        "sections_missing": missing_sections,
        "information_required_markers": info_required,
        "sections_finalized": finalized,
        "ready_for_export": not missing_sections and info_required == 0,
    }


def _markdown(req: ExportRequest) -> str:
    lines = [f"# Draft Red Herring Prospectus — {req.company_name}",
             f"\n*Generated {date.today().isoformat()} · DRAFT — not for regulatory submission*\n"]
    for s in req.sections:
        lines.append(s.content.strip() + "\n")
    return "\n".join(lines)


def _html(md_company: str, req: ExportRequest) -> str:
    body = []
    for s in req.sections:
        # minimal markdown -> html for headings/paragraphs
        html_content = []
        for para in s.content.strip().split("\n\n"):
            para = para.strip()
            if para.startswith("## "):
                html_content.append(f"<h2>{para[3:].strip()}</h2>")
            elif para.startswith("*") and para.endswith("*"):
                html_content.append(f"<p><em>{para.strip('*').strip()}</em></p>")
            else:
                html_content.append(f"<p>{para.replace(chr(10), '<br>')}</p>")
        body.append("\n".join(html_content))
    joined = "\n".join(body)
    return f"""<!doctype html><html><head><meta charset="utf-8">
<title>DRHP — {md_company}</title>
<style>@page{{size:A4;margin:2cm}}body{{font-family:Georgia,serif;line-height:1.5;max-width:800px;margin:2rem auto}}
h1{{border-bottom:2px solid #333}}h2{{margin-top:2rem;color:#1a3a5c}}em{{color:#666}}
.banner{{background:#fff3cd;border:1px solid #ffc107;padding:.5rem;font-size:.9rem}}</style></head>
<body><h1>Draft Red Herring Prospectus — {md_company}</h1>
<p class="banner">DRAFT — subject to mandatory review and certification by authorised
intermediaries before any submission to SEBI.</p>
{joined}</body></html>"""


def _docx(req: ExportRequest) -> str | None:
    try:
        import docx  # python-docx
    except ImportError:
        return None
    document = docx.Document()
    document.add_heading(f"Draft Red Herring Prospectus — {req.company_name}", level=0)
    document.add_paragraph("DRAFT — not for regulatory submission").italic = True
    for s in req.sections:
        for para in s.content.strip().split("\n\n"):
            para = para.strip()
            if para.startswith("## "):
                document.add_heading(para[3:].strip(), level=1)
            elif para.startswith("*") and para.endswith("*"):
                document.add_paragraph(para.strip("*").strip())
            else:
                document.add_paragraph(para)
    buf = io.BytesIO()
    document.save(buf)
    return base64.b64encode(buf.getvalue()).decode("ascii")


@router.post("/export")
def export(req: ExportRequest) -> dict:
    docx_b64 = _docx(req)
    return {
        "company_name": req.company_name,
        "markdown": _markdown(req),
        "html": _html(req.company_name, req),
        "docx_base64": docx_b64,
        "docx_available": docx_b64 is not None,
        "summary": _compliance_summary(req),
    }
